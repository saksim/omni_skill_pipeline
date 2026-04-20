from __future__ import annotations

import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from typing import Dict, List, Protocol, Sequence

from omni_skill_pipeline.extraction.modality.document_parser import DocumentStructureParser
from omni_skill_pipeline.models import Asset, ContentType, EvidenceUnit, LoadedAsset, Modality, TextDistillRequest
from omni_skill_pipeline.utils import read_text_file, split_paragraphs, unique_preserve_order


class TextReader(Protocol):
    suffixes: Sequence[str]

    def read(self, path: Path) -> str:
        ...


class PlainTextReader(object):
    suffixes = ('.txt', '.md', '.markdown', '.rst', '.log')

    def read(self, path: Path) -> str:
        return read_text_file(path)


class JsonTextReader(object):
    suffixes = ('.json',)

    def read(self, path: Path) -> str:
        data = json.loads(path.read_text(encoding='utf-8'))
        return json.dumps(data, ensure_ascii=False, indent=2)


class HtmlTextReader(object):
    suffixes = ('.html', '.htm')

    def read(self, path: Path) -> str:
        raw_html = read_text_file(path)
        try:
            from bs4 import BeautifulSoup  # type: ignore
        except ImportError:
            return re.sub(r'<[^>]+>', ' ', raw_html)
        soup = BeautifulSoup(raw_html, 'html.parser')
        return soup.get_text('\n')


class DocxReader(object):
    suffixes = ('.docx',)

    def read(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return self._read_via_zip(path)
        document = Document(str(path))
        lines: list[str] = []
        code_buffer: list[str] = []

        def flush_code_buffer() -> None:
            if not code_buffer:
                return
            lines.append('```')
            lines.extend(code_buffer)
            lines.append('```')
            lines.append('')
            code_buffer.clear()

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                flush_code_buffer()
                lines.append('')
                continue
            style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").strip()
            lowered_style = style_name.lower()

            if "code" in lowered_style:
                code_buffer.append(text)
                continue

            flush_code_buffer()
            heading_level = self._heading_level_from_style(style_name)
            if heading_level is not None:
                lines.append("%s %s" % ("#" * heading_level, text))
                lines.append('')
                continue

            toc_level = self._toc_level_from_style(style_name)
            if toc_level is not None:
                lines.append("[TOC L%s] %s" % (toc_level, text))
                continue

            lines.append(text)
            lines.append('')

        flush_code_buffer()

        for table in document.tables:
            markdown_table = self._table_to_markdown(table)
            if markdown_table:
                if lines and lines[-1] != '':
                    lines.append('')
                lines.extend(markdown_table)
                lines.append('')

        return "\n".join(lines).strip()

    def _heading_level_from_style(self, style_name: str) -> int | None:
        lowered = style_name.lower()
        if not lowered.startswith("heading"):
            return None
        match = re.search(r"(\d+)", lowered)
        if match is None:
            return 1
        try:
            level = int(match.group(1))
        except ValueError:
            return 1
        return max(1, min(level, 6))

    def _toc_level_from_style(self, style_name: str) -> int | None:
        lowered = style_name.lower()
        if "toc" not in lowered and "contents" not in lowered:
            return None
        match = re.search(r"(\d+)", lowered)
        if match is None:
            return 1
        try:
            return max(1, int(match.group(1)))
        except ValueError:
            return 1

    def _table_to_markdown(self, table) -> list[str]:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cell for cell in cells):
                rows.append(cells)
        if not rows:
            return []
        width = max(len(row) for row in rows)
        normalized_rows = [row + [""] * (width - len(row)) for row in rows]
        lines = ["| " + " | ".join(normalized_rows[0]) + " |"]
        lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in normalized_rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return lines

    def _read_via_zip(self, path: Path) -> str:
        with zipfile.ZipFile(str(path)) as archive:
            xml = archive.read('word/document.xml').decode('utf-8')
        normalized = xml.replace('</w:p>', '\n\n').replace('<w:tab/>', ' ')
        cleaned = re.sub(r'<[^>]+>', '', normalized)
        return unescape(cleaned)


class PdfReader(object):
    suffixes = ('.pdf',)

    def read(self, path: Path) -> str:
        try:
            from pypdf import PdfReader as NativePdfReader  # type: ignore
        except ImportError:
            return self._read_via_pdftotext(path)
        reader = NativePdfReader(str(path))
        pages = [page.extract_text() or '' for page in reader.pages]
        text = '\n\n'.join(page.strip() for page in pages if page.strip())
        if text:
            return text
        return self._read_via_pdftotext(path)

    def _read_via_pdftotext(self, path: Path) -> str:
        executable = shutil.which('pdftotext')
        if not executable:
            raise ValueError('PDF support requires pypdf or pdftotext.')
        process = subprocess.run(
            [executable, '-layout', str(path), '-'],
            check=False,
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore',
        )
        if process.returncode != 0:
            raise ValueError(process.stderr.strip() or 'pdftotext failed.')
        return process.stdout


class DocReader(object):
    suffixes = ('.doc',)

    def read(self, path: Path) -> str:
        executable = shutil.which('antiword')
        if not executable:
            raise ValueError('DOC support requires antiword.')
        process = subprocess.run(
            [executable, str(path)],
            check=False,
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore',
        )
        if process.returncode != 0:
            raise ValueError(process.stderr.strip() or 'antiword failed.')
        return process.stdout


@dataclass(slots=True)
class TextReaderRegistry:
    readers: Sequence[TextReader]

    def get_reader(self, suffix: str) -> TextReader:
        normalized = suffix.lower()
        for reader in self.readers:
            if normalized in reader.suffixes:
                return reader
        raise ValueError('Unsupported text format: %s' % normalized)


class TextAdapter(object):
    def __init__(
        self,
        registry: TextReaderRegistry | None = None,
        document_parser: DocumentStructureParser | None = None,
    ) -> None:
        self.registry = registry or TextReaderRegistry(
            readers=[
                PlainTextReader(),
                JsonTextReader(),
                HtmlTextReader(),
                DocxReader(),
                PdfReader(),
                DocReader(),
            ]
        )
        self.document_parser = document_parser or DocumentStructureParser()

    def load(self, request: TextDistillRequest) -> LoadedAsset:
        request.validate()
        if request.content:
            raw_text = request.content
            source_uri = request.file_path or 'inline://text'
            filename = Path(request.file_path).name if request.file_path else 'inline.txt'
        else:
            path = Path(str(request.file_path))
            raw_text = self._read_path(path)
            source_uri = str(path.resolve())
            filename = path.name

        title_hint = self._derive_title_hint(request.title, raw_text, filename)
        asset = Asset(
            modality=Modality.TEXT,
            source_uri=source_uri,
            metadata={'filename': filename, 'language': 'mixed'},
        )
        source_format = Path(filename).suffix.lower() or 'inline'
        evidence_units = self._build_evidence(asset.asset_id, raw_text, source_format=source_format)
        return LoadedAsset(
            asset=asset,
            evidence_units=evidence_units,
            title_hint=title_hint,
            adapter_metadata={'format': source_format},
        )

    def _read_path(self, path: Path) -> str:
        return self.registry.get_reader(path.suffix).read(path)

    def _derive_title_hint(self, explicit_title: str | None, raw_text: str, filename: str) -> str:
        if explicit_title:
            return explicit_title.strip()
        for line in raw_text.splitlines():
            stripped = line.strip()
            if stripped.startswith('#'):
                return stripped.lstrip('#').strip()
            if stripped:
                return stripped[:80]
        return Path(filename).stem.replace('_', ' ')

    def _build_evidence(self, asset_id: str, raw_text: str, *, source_format: str = "") -> List[EvidenceUnit]:
        parsed_blocks = self.document_parser.parse(raw_text, source_format=source_format)
        if parsed_blocks:
            evidence_units = []
            for block in parsed_blocks:
                evidence_units.append(
                    EvidenceUnit(
                        asset_id=asset_id,
                        span_ref=block.span_ref,
                        content_type=block.content_type,
                        content=block.content,
                        confidence=block.confidence,
                        tags=unique_preserve_order(block.tags),
                    )
                )
            return evidence_units

        paragraphs = split_paragraphs(raw_text)
        evidence_units = []
        for index, paragraph in enumerate(paragraphs, start=1):
            tags = []
            lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
            if lines and lines[0].startswith('#'):
                tags.append('heading')
            evidence_units.append(
                EvidenceUnit(
                    asset_id=asset_id,
                    span_ref='paragraph:%04d' % index,
                    content_type=ContentType.TEXT,
                    content=paragraph,
                    confidence=0.86,
                    tags=unique_preserve_order(tags),
                )
            )
        if not evidence_units and raw_text.strip():
            evidence_units.append(
                EvidenceUnit(
                    asset_id=asset_id,
                    span_ref='paragraph:0001',
                    content_type=ContentType.TEXT,
                    content=raw_text.strip(),
                    confidence=0.72,
                    tags=[],
                )
            )
        return evidence_units
